import os
from pathlib import Path
from typing import Dict, Any, Optional

from llama_index.core import VectorStoreIndex, StorageContext, Document
from qdrant_client import QdrantClient
from llama_index.vector_stores.qdrant import QdrantVectorStore

from .settings import configure_llamaindex, qdrant_config

from pypdf import PdfReader
import docx

DOCS_DIR = Path(os.getenv("DOCS_DIR", "/docs/sample"))


def _parse_file(file_path: str) -> str:
    p = Path(file_path)
    ext = p.suffix.lower()

    if ext in [".txt", ".md", ".markdown"]:
        return p.read_text(encoding="utf-8")

    if ext == ".pdf":
        reader = PdfReader(str(p))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        return "\n".join(texts)

    if ext == ".docx":
        d = docx.Document(str(p))
        return "\n".join([para.text for para in d.paragraphs if para.text])

    raise ValueError(f"unsupported_file_type: {ext}")


def _build_index(client: QdrantClient, collection: str):
    vector_store = QdrantVectorStore(client=client, collection_name=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    return storage_context


def ingest_dir(collection: Optional[str] = None) -> dict:
    configure_llamaindex()
    url, default_collection = qdrant_config()
    col = collection or default_collection
    client = QdrantClient(url=url)

    # Simple directory reader: txt/md 중심
    docs: list[Document] = []
    for fp in DOCS_DIR.rglob("*"):
        if fp.is_file() and fp.suffix.lower() in [".md", ".txt", ".markdown"]:
            docs.append(Document(text=fp.read_text(encoding="utf-8"), metadata={"source": str(fp.relative_to(DOCS_DIR))}))

    storage_context = _build_index(client, col)
    _ = VectorStoreIndex.from_documents(docs, storage_context=storage_context)

    return {"inserted_docs": len(docs), "collection": col}


def ingest_file(file_path: str, collection: str, metadata: Optional[Dict[str, Any]] = None) -> dict:
    configure_llamaindex()
    url, _ = qdrant_config()
    client = QdrantClient(url=url)

    text = _parse_file(file_path)
    md = metadata or {}
    md.setdefault("source", Path(file_path).name)
    md.setdefault("file", file_path)

    doc = Document(text=text, metadata=md)

    storage_context = _build_index(client, collection)
    _ = VectorStoreIndex.from_documents([doc], storage_context=storage_context)

    return {"inserted_docs": 1, "collection": collection, "source": md.get("source")}
