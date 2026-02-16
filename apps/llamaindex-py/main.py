import os
import fitz  # PyMuPDF
import docx
from fastapi import FastAPI
from pydantic import BaseModel
from qdrant_client import QdrantClient
from qdrant_client.http.models import VectorParams, Distance
from llama_index.core import Document, VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.qdrant import QdrantVectorStore

from opentelemetry.instrumentation.fastapi import FastAPIInstrumentor
from opentelemetry.instrumentation.requests import RequestsInstrumentor

from otel.python.otel import setup_otel

setup_otel("llamaindex-py")

app = FastAPI(title="llamaindex-py")
FastAPIInstrumentor.instrument_app(app)
RequestsInstrumentor().instrument()

QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

def qdrant_client():
    return QdrantClient(url=QDRANT_URL)

def provider():
    return os.getenv("LLM_PROVIDER", "ollama").lower()

def embed_model():
    if provider() == "ollama":
        from llama_index.embeddings.ollama import OllamaEmbedding
        return OllamaEmbedding(
            model_name=os.getenv("OLLAMA_EMBED_MODEL", "nomic-embed-text"),
            base_url=os.getenv("OLLAMA_BASE_URL", "http://ollama:11434"),
        )
    from llama_index.embeddings.openai import OpenAIEmbedding
    return OpenAIEmbedding(
        model=os.getenv("OPENAI_EMBEDDING_MODEL", "text-embedding-3-large"),
        api_key=os.getenv("OPENAI_API_KEY"),
    )

def llm_model():
    if provider() == "ollama":
        from llama_index.llms.ollama import Ollama
        return Ollama(
            model=os.getenv("OLLAMA_LLM_MODEL", "llama3.1:8b"),
            base_url=os.getenv("OLLAMA_BASE_URL", "http://ollama:11434"),
            temperature=0.2,
        )
    from llama_index.llms.openai import OpenAI
    return OpenAI(
        model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        api_key=os.getenv("OPENAI_API_KEY"),
        temperature=0.2,
    )

class IngestReq(BaseModel):
    collection: str
    filePath: str
    fileName: str | None = None
    mimeType: str | None = None
    docId: str | None = None
    version: int | None = None
    sha256: str | None = None
    tenant: str | None = None
    dept: str | None = None

class ChatReq(BaseModel):
    question: str
    collection: str
    topK: int = 5

def read_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".md", ".txt"]:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".pdf":
        doc = fitz.open(file_path)
        text = []
        for page in doc:
            text.append(page.get_text("text"))
        out = "\n".join(text).strip()
        return out or "[PDF에서 추출된 텍스트가 없습니다. (스캔 PDF일 수 있음)]"

    if ext == ".docx":
        d = docx.Document(file_path)
        out = "\n".join([p.text for p in d.paragraphs]).strip()
        return out or "[DOCX에서 추출된 텍스트가 없습니다.]"

    return f"지원되지 않는 파일 형식({ext})입니다."

def ensure_collection(name: str, vector_size: int):
    client = qdrant_client()
    collections = client.get_collections().collections
    if not any(c.name == name for c in collections):
        client.create_collection(
            collection_name=name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )
        return
    info = client.get_collection(name)
    existing = getattr(getattr(info.config.params, "vectors", None), "size", None)
    if existing and int(existing) != int(vector_size):
        raise ValueError(
            f"Qdrant collection '{name}' vector size mismatch: existing={existing}, requested={vector_size}. "
            "DEV에서는 컬렉션 삭제 또는 tenant/dept 변경을 권장합니다."
        )

@app.post("/ingest/file")
def ingest_file(req: IngestReq):
    text = read_text(req.filePath)
    metadata = {
        "source": req.fileName or req.filePath,
        "docId": req.docId,
        "version": req.version,
        "sha256": req.sha256,
        "tenant": req.tenant,
        "dept": req.dept,
        "embedModel": os.getenv("OLLAMA_EMBED_MODEL") if provider() == "ollama" else os.getenv("OPENAI_EMBEDDING_MODEL"),
    }
    docs = [Document(text=text, metadata=metadata)]

    splitter = SentenceSplitter(chunk_size=800, chunk_overlap=120)
    nodes = splitter.get_nodes_from_documents(docs)

    emb = embed_model()
    # infer dimension by embedding one string
    sample = emb.get_text_embedding("vector-size-check")
    ensure_collection(req.collection, len(sample))

    vs = QdrantVectorStore(client=qdrant_client(), collection_name=req.collection)
    index = VectorStoreIndex(nodes, vector_store=vs, embed_model=emb)

    _ = index.as_query_engine(llm=llm_model())
    return {"ok": True, "collection": req.collection, "chunks": len(nodes), "vectorSize": len(sample)}

@app.post("/chat")
def chat(req: ChatReq):
    vs = QdrantVectorStore(client=qdrant_client(), collection_name=req.collection)
    index = VectorStoreIndex.from_vector_store(vector_store=vs, embed_model=embed_model())

    qe = index.as_query_engine(similarity_top_k=req.topK, llm=llm_model())
    resp = qe.query(
        "너는 은행업무 보조 AI다. 컨텍스트에 없는 내용은 추측하지 말고 부족하다고 말해. "
        "준법/보안 관점에서 고객에게 노출하면 안 되는 정보를 생성하지 않는다.\n"
        f"질문: {req.question}"
    )
    return {"answer": str(resp), "collection": req.collection}

@app.get("/healthz")
def healthz():
    return {"ok": True, "provider": provider()}
